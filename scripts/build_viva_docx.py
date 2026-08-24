#!/usr/bin/env python3
"""Build the internal viva/project Word guide. Not committed to GitHub."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from pathlib import Path

OUT = Path(__file__).resolve().parents[1] / "AquaIndex-Project-Viva-Guide.docx"

QA = [
    (
        "What is this project in one minute?",
        "AquaIndex is a web dashboard that evaluates groundwater using Digital Water Level Recorder (DWLR) style time series. "
        "It shows an India map of station status, charts of water level in metres below ground, estimates monsoon recharge with the water-table fluctuation method, "
        "computes a 0–100 availability index, and lets a planner try a simple extraction-cut scenario. "
        "It is a Computer Science prototype for CSE 2025–26, not an official CGWB operations system.",
    ),
    (
        "What problem are you solving?",
        "Groundwater is over-used in many Indian districts, but much monitoring is still seasonal and paper-based. "
        "High-frequency DWLR data exists nationally (about 5,260 recorders in our synopsis) but is hard for a local planner to use. "
        "We convert piezometer levels into recharge, health status, and a map a non-modeller can click through.",
    ),
    (
        "Expand DWLR / DWLR.",
        "Digital Water Level Recorder (also written DWLR). A logger in an observation well that records water level automatically, often hourly or daily, instead of a person dropping a tape once a month.",
    ),
    (
        "Expand CGWB, mbgl, Sy, WTF.",
        "CGWB: Central Ground Water Board, Ministry of Jal Shakti. "
        "mbgl: metres below ground level (larger number = deeper, more stressed water table). "
        "Sy: specific yield, drainable porosity of the aquifer (typical alluvium 0.10–0.16, hard rock 0.015–0.03). "
        "WTF: water-table fluctuation method for recharge.",
    ),
    (
        "What are CGWB assessment classes?",
        "Safe, Semi-Critical, Critical, Over-Exploited (our UI uses these four). They come from comparing annual extraction with annual extractable resource at assessment-unit scale. "
        "Our stations are tagged with the class that is typical of that district type in published compilations. We do not recompute the official CGWB GIS layer.",
    ),
    (
        "Is this real-time government data?",
        "No. Be honest. India-WRIS / NWIC / CGWB telemetry is not a free public JSON API we can call from a student website without access approval. "
        "We use a representative subset of 65 sites with realistic coordinates and classes, plus generated daily series that behave like DWLR records (monsoon rise, pumping decline). "
        "The live tick is a UI heartbeat, not a satellite link to a well.",
    ),
    (
        "Then how is the data 'obtained'?",
        "Station metadata (name, district, lat, lng, aquifer, Sy, CGWB class, stress) is authored in src/data/stations.js from public geography and published stress patterns. "
        "Daily levels for ~92 days (late May to 24 August 2026 in the demo clock) are simulated with a seeded PRNG (mulberry32) so every student sees the same series. "
        "Rainfall on the chart is a monsoon proxy used inside that generator, not IMD AWS data.",
    ),
    (
        "Why 65 stations if the synopsis says 5,260?",
        "5,260 is the national DWLR network size we cite as context. Loading thousands of series in a free static website would be slow and would claim data we do not legally stream. "
        "65 sites cover major states and both alluvial and hard-rock aquifers so the map, recharge table, and planner still demonstrate the full workflow.",
    ),
    (
        "Explain: Live tick 22 · demo subset 65 of 5,260 stations.",
        "Live tick 22: Overview page counter. Starts at 0, +1 every 7 seconds. Each tick adds a sine jitter of about ±1 cm to the displayed latest mbgl so KPIs look alive. "
        "Demo subset 65: stations actually in the app. Of 5,260: real national network size from the project write-up. Say this sentence if asked: we prototype the information system on a subset; we do not host the full national feed.",
    ),
    (
        "Walk me through a live demo.",
        "1) Overview: four KPIs, map colours, lowest-index list. 2) DWLR map: filter Karnataka or Punjab, click a circle, Open station. "
        "3) Station: inverted mbgl chart, rainfall bars, WTF mm, index. 4) Recharge: formula and table. "
        "5) Planner: Bengaluru Anekal, move extraction cut, show metres of depth avoided. 6) Project: team, Prof. Prachitha M., disclaimer.",
    ),
    (
        "What does each menu item do?",
        "Overview / — national picture. DWLR map /map — filter and Leaflet map. Station /station/:id — one well. Recharge /recharge — WTF ranking. Planner /planner — extraction scenario. Project /about — team and references.",
    ),
    (
        "Why is the water-level axis inverted?",
        "Hydrologists plot mbgl so a rising water table (improvement) appears as an upward line. We set Recharts YAxis reversed. Always say: 25 mbgl is deeper than 10 mbgl.",
    ),
    (
        "Formula for recharge.",
        "R (mm) = sum of (Δh × Sy × 1000) over days when the water table rises (mbgl decreases) by more than 0.008 m. "
        "Δh is in metres. Times 1000 converts metres of water in the pores to millimetres of recharge equivalent. Implemented in enrich() in stations.js.",
    ),
    (
        "Why not use rainfall × infiltration coefficient only?",
        "Rainfall-recharge factors are useful at basin scale but ignore the actual piezometer. WTF uses the well itself: if the table rose, that rise times Sy is a local recharge indicator. We still show a rainfall proxy on the chart for teaching.",
    ),
    (
        "What is the availability index?",
        "Score 0–100 = 45% depth score + 30% 90-day trend score + 25% recharge score. "
        "≥70 Good, ≥50 Moderate, ≥30 Stressed, else Poor. It is our composite for the public, not an official CGWB index.",
    ),
    (
        "What does the planner actually compute?",
        "dailySlope = delta90 / 90. Future mbgl ≈ current + dailySlope × (1 − cut%) × 365 − extraRechargeTerm × 365. "
        "Cut% is reduced pumping. Extra term is a small bonus for leaving more monsoon water in the aquifer. Screening only — not MODFLOW.",
    ),
    (
        "What is unique / innovative?",
        "1) One lightweight site from map to recharge to health to intervention. 2) WTF on daily-style DWLR traces, not only annual CGWB PDFs. "
        "3) Availability index combining depth, trend, recharge. 4) Free static hosting instead of AWS. 5) Transparent subset + live-tick labelling so we do not fake a classified government login.",
    ),
    (
        "Which APIs do you use, and whose are they?",
        "Map tiles: CARTO Dark Matter tiles at basemaps.cartocdn.com, map data © OpenStreetMap contributors — no API key. "
        "Optional LLM: Google Gemini 2.0 Flash, Generative Language API, generativelanguage.googleapis.com, env VITE_GEMINI_API_KEY. "
        "We do not call CGWB, India-WRIS, Google Maps, Mapbox, or AWS APIs. There is no student-written backend REST server in production.",
    ),
    (
        "If Gemini is down?",
        "briefStation() always has localBrief() from the same numbers. The hydrologist paragraph still appears. The product does not depend on Google.",
    ),
    (
        "Why frontend-only? The synopsis said Flask/Django and AWS.",
        "The teacher needed a working demo quickly, with zero hosting cost. A static React app on Vercel meets that. "
        "The 'API structure' is the station JSON model in the browser. A future phase can put the same JSON behind FastAPI and swap the generator for NWIC once access exists. Do not claim we already deployed AWS.",
    ),
    (
        "Tech stack?",
        "React 19, Vite 5, React Router 7, Tailwind CSS 4, Leaflet, Recharts, optional Gemini. Node.js to build. Host: static files (HTML/JS/CSS).",
    ),
    (
        "Why Leaflet not Mapbox?",
        "Mapbox tokens are paid after a quota. Leaflet + CARTO/OSM tiles are free and enough for a national station map.",
    ),
    (
        "How do you keep 65 series identical for every visitor?",
        "mulberry32 seeded PRNG: seed = 18000 + index×97. Same code, same series, no database.",
    ),
    (
        "Security: is the Gemini key in GitHub?",
        "It must not be. .env is gitignored. Only VITE_GEMINI_API_KEY in Vercel env. Keys in chat should be rotated. A VITE_ key is still visible in the browser bundle — for a class demo only; production should proxy the LLM.",
    ),
    (
        "SPA routing on Vercel?",
        "vercel.json rewrites all paths to index.html so /map and /station/DWLR-KA-001 do not 404 on refresh.",
    ),
    (
        "Team roles vs who wrote the website?",
        "Official roles: Abhishek Biswal (estimation logic), Shivam Kumar (GIS/viz), Sushant Kumar (API/data layer), Srushti S Mopagar (UI/architecture). Guide: Prof. Prachitha M. "
        "If asked who coded the React app, say the team implemented the prototype together; do not invent a company.",
    ),
    (
        "Bengaluru Anekal — why is it Over-Exploited?",
        "Bengaluru Urban is a classic hard-rock, low-Sy, high-pumping setting. Deep mbgl, weak monsoon recovery in the demo series, high stress parameter. Good planner demo.",
    ),
    (
        "Punjab vs Assam on your map?",
        "Punjab alluvium, intensive paddy, Over-Exploited / Critical, deep mbgl. Assam Brahmaputra alluvium, high Sy, Safe, shallow table. Shows the index is not one colour for all of India.",
    ),
    (
        "Hard rock vs alluvium for recharge millimetres.",
        "Same 0.5 m rise: alluvium Sy 0.12 → 60 mm; gneiss Sy 0.015 → 15 mm. That is why Kolar may show small mm even if the chart wiggles.",
    ),
    (
        "Limitations?",
        "No live telemetry; 65 not 5,260; rainfall proxy; planner not calibrated; Gemini not official advice; no quality flags, well construction, or pumping tests.",
    ),
    (
        "Future work?",
        "NWIC/CGWB authenticated API or bulk CSV; district GeoJSON choropleth; more than daily (hourly) DWLR; couple IMD rainfall; replace planner with a lumped reservoir or MODFLOW; PWA for field staff.",
    ),
    (
        "Is this machine learning / AI project?",
        "Core is hydrology + visualisation. Optional Gemini is NLP wrapping of numbers we already computed — not a predicted water table. Do not call it a deep-learning recharge model.",
    ),
    (
        "Testing?",
        "npm run build must succeed. Manually test map filters, station 404, planner slider, recharge sort. No Cypress suite in this delivery — say we used build + walkthrough due to the demo deadline.",
    ),
    (
        "Why dark UI?",
        "Operations-console metaphor (live monitoring). Teal for water, red/amber for stress. Not required scientifically.",
    ),
    (
        "References?",
        "[1] CGWB National Compilation on Dynamic Groundwater Resources of India, MoJS, 2024. [2] Nicholson et al., Journal of Hydrology, 2023, real-time groundwater monitoring. [3] IEEE Std 2510-2021 sensor data middleware. Cite as inspiration, not as 'we implemented the whole IEEE stack'.",
    ),
    (
        "Difference between observation well and pumping well?",
        "DWLR in the demo is treated as a piezometer (observation). Pumping wells have drawdown cones. WTF recharge on a pumped well can be biased. We assume monitoring wells.",
    ),
    (
        "What is specific yield vs porosity vs storativity?",
        "Porosity is all voids. Specific yield Sy is the volume that drains by gravity (unconfined aquifers) — what WTF uses. Storativity applies more to confined aquifers and is much smaller. We use Sy only.",
    ),
    (
        "Could you show over-exploitation scientifically?",
        "Officially, extraction / extractable resource > 100% at assessment unit. We do not have extraction volumes. We use class tags plus falling mbgl as a consistent story, and say a full budget needs draft data from CGWB/state GWWs.",
    ),
    (
        "How would you ingest real DWLR CSV tomorrow?",
        "Keep station id, timestamp, mbgl. Replace buildSeries() with parse CSV → same enrich() for WTF and index. Map component already only needs lat, lng, cat, latest. That is the intended architecture.",
    ),
    (
        "CORS and Gemini from the browser?",
        "The Generative Language API is called with fetch from the client. If CORS or key errors occur, we catch and use localBrief. That is why the page always fills.",
    ),
    (
        "Performance?",
        "65 × 92 points is tiny. Leaflet circle markers, Recharts on one station, lazy() on map and pages, Carto tiles cached by the browser. Bundle split so Leaflet is not required to paint the about page.",
    ),
    (
        "Ethical / policy angle?",
        "Showing Over-Exploited in red can alarm people. We label it a prototype. Groundwater data can be sensitive for drinking-water security; we do not scrape restricted portals.",
    ),
    (
        "What if the teacher opens /map and refreshes after deploy?",
        "Without SPA rewrites the host looks for a file /map and 404s. vercel.json and public/_redirects send all routes to index.html. React Router then renders MapPage.",
    ),
    (
        "Unit of recharge vs unit of level?",
        "Level: metres below ground. Recharge: millimetres of water added to the aquifer column (like rainfall units) so it can be compared with monsoon mm.",
    ),
    (
        "Why August 2026 in the series?",
        "The demo clock is aligned with the academic year and a monsoon window (May–August) so rise events exist for WTF. It is the narrative date of the prototype dataset.",
    ),
    (
        "Can this run without internet?",
        "After npm run build, the JS/CSS work offline, but the map tiles need the network (Carto). Charts and tables still work if Leaflet tiles fail (grey map).",
    ),
    (
        "How is 'falling tables' counted?",
        "Stations with delta90 > 0.15 m (water table deeper by more than 15 cm over ~90 days). Rising: delta90 < −0.15 m.",
    ),
    (
        "What should we not say?",
        "Do not say we hacked CGWB. Do not say Gemini predicts groundwater. Do not say the planner is government-approved. Do not say all 5,260 wells are in the database. Do not paste API keys on slides.",
    ),
    (
        "One closing sentence if they ask 'so what?'",
        "If a block officer can see that Anekal is over-exploited, that WTF recharge this monsoon is only a few tens of millimetres in hard rock, and that cutting pumping slows further decline, they have a starting point for demand management — that is the point of the prototype.",
    ),
    (
        "Where is the source code?",
        "GitHub (public repository after this submission). README explains run and deploy. The Word viva guide is local only and not in git.",
    ),
]


def set_run_font(run, size=11, bold=False, color=None, name="Calibri"):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = name
    r = run._element.get_or_add_rPr()
    rFonts = r.get_or_add_rFonts()
    rFonts.set(qn("w:eastAsia"), name)
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_para(doc, text, size=11, bold=False, space_after=8, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def main():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.85)
        s.bottom_margin = Inches(0.85)
        s.left_margin = Inches(1.0)
        s.right_margin = Inches(1.0)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("AquaIndex")
    set_run_font(r, size=26, bold=True, color=(15, 42, 40))

    st = doc.add_paragraph()
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = st.add_run(
        "Real-Time Groundwater Resource Evaluation Using DWLR Data\n"
        "Project explanation, architecture, APIs, demo script, and 50 viva questions\n"
        "B.E. / B.Tech · Computer Science and Engineering · 2025–26"
    )
    set_run_font(r, size=12, color=(45, 70, 68))

    add_para(
        doc,
        "Guide: Prof. Prachitha M.   |   Team: Abhishek Biswal · Shivam Kumar · Sushant Kumar · Srushti S Mopagar",
        size=11,
        bold=True,
    )
    add_para(
        doc,
        "This document is for the team and the oral exam. It is not uploaded to GitHub. Read it once end-to-end before the demo.",
        size=11,
    )

    doc.add_heading("1. What this project is (plain language)", level=1)
    add_para(
        doc,
        "Imagine thousands of electronic rulers hanging in wells across India. Each ruler (a DWLR) writes down how far the water is below the ground, every day or more often. "
        "Scientists in CGWB already collect much of this. Ordinary planners and students cannot easily open that firehose and see: is my district recovering this monsoon, and if we pump less, does it help?",
    )
    add_para(
        doc,
        "AquaIndex is a website that tells that story. You open a map of India. Dots are observation wells. Colour is CGWB-style stress class. Click a well: you see the last ~90 days of water level, a rainfall-like bar, how many millimetres of recharge we estimate, and a health score. "
        "A planner page asks: if this area cut pumping by 20%, how much less deep might the water table be in a year? That last number is a classroom model, not a court-ready simulation.",
    )

    doc.add_heading("2. How to speak in the viva (attitude)", level=1)
    add_para(
        doc,
        "Lead with the hydrology workflow, then the software. If they ask about live CGWB APIs, tell the truth immediately (section 5). Teachers respect a bounded prototype more than a fake 'we connected to Jal Shakti'. "
        "Use the words metres below ground, specific yield, water-table fluctuation, Over-Exploited. Do not say 'the AI predicts groundwater' unless you immediately correct: Gemini only writes four sentences from numbers we already calculated.",
    )

    doc.add_heading("3. Demo script in front of the teacher (about four minutes)", level=1)
    steps = [
        "Open the live URL. Stay on Overview. Point to mean mbgl and falling vs recovering counts. Mention live tick is a display heartbeat.",
        "Gesture at the map: red Over-Exploited (Punjab, parts of Karnataka), green Safe (northeast, much of Kerala).",
        "Click a name under lowest availability. You land on a station page.",
        "On the station chart: 'Axis is inverted; water table rising means the line goes up, mbgl number goes down.' Mention WTF millimetres and Sy.",
        "Go to DWLR map. Filter state Karnataka, class Over-Exploited. Click Bengaluru Anekal.",
        "Go to Recharge. Read the formula R = Σ Δh × Sy × 1000. Compare an alluvial Punjab site vs a hard-rock Karnataka site.",
        "Planner: Anekal, slide extraction cut to ~30%. 'Depth avoided' is screening, not MODFLOW.",
        "Project page: team, guide, CGWB 2024 / Hydrology 2023 / IEEE 2510 references, disclaimer that series are demonstration traces.",
    ]
    for i, s in enumerate(steps, 1):
        add_para(doc, f"{i}. {s}", space_after=4)

    doc.add_heading("4. Screens and what every number means", level=1)
    add_para(doc, "Overview KPIs", bold=True, space_after=4)
    add_para(
        doc,
        "National index: average of 65 station availability scores. Mean water level: average latest mbgl. Falling tables: count with 90-day deepening > 0.15 m. Mean WTF recharge: average millimetres over the demo monsoon window.",
    )
    add_para(doc, "Station KPIs", bold=True, space_after=4)
    add_para(
        doc,
        "Latest level (mbgl). 90-day change (positive = deeper/worse). WTF recharge (mm) and count of rise events. Availability index and health word. Briefing paragraph: Gemini or local template.",
    )
    add_para(doc, "Planner", bold=True, space_after=4)
    add_para(
        doc,
        "Now / 1 year no change / 1 year with cut. Depth avoided = baseline projected mbgl minus cut scenario (positive means we avoided going that much deeper).",
    )

    doc.add_heading("5. Data: what is real, what is generated, whose APIs", level=1)
    add_para(doc, "5.1 Not used (say this clearly)", bold=True, space_after=4)
    add_para(
        doc,
        "No REST calls to CGWB, India-WRIS, NWIC, Bhuvan groundwater layers, or IMD as a live service. No AWS. No Mapbox token. No database.",
    )
    add_para(doc, "5.2 What we did use", bold=True, space_after=4)
    add_para(
        doc,
        "Authored metadata for 65 DWLR-like sites: real districts, plausible lat/lng, aquifer labels, Sy, CGWB class consistent with known regional stress (e.g. Sangrur Over-Exploited, Jorhat Safe). "
        "Time series: generated in JavaScript with a fixed seed; monsoon months get more 'rain' proxy and more chance of water-table rise; high stress sites keep declining. "
        "Map pictures: CARTO CDN tiles (basemaps.cartocdn.com/dark_all), cartography by CARTO, geographic data © OpenStreetMap contributors. This is an HTTP tile API in the loose sense (z/x/y.png), no key, not a groundwater API. "
        "Optional text: Google Gemini API (Google AI Studio / Generative Language API), model gemini-2.0-flash, browser fetch to generativelanguage.googleapis.com. Key in VITE_GEMINI_API_KEY. Fallback always exists.",
    )
    add_para(doc, "5.3 Why this is still a valid CSE project", bold=True, space_after=4)
    add_para(
        doc,
        "The software architecture is ready for real CSV/API: same station object, same WTF and index functions. The contribution is the interactive evaluation pipeline and the hydrology methods implemented in code, plus a deployable UI. "
        "Hydrology projects often use synthetic or reconstructed series when operational feeds are closed; we label that limitation on the Project page.",
    )

    doc.add_heading("6. Algorithms (enough to write on the board)", level=1)
    add_para(
        doc,
        "Water-table fluctuation recharge. For each consecutive pair of daily levels L[i-1], L[i] (mbgl): if L[i-1] − L[i] > 0.008, that difference is a rise Δh in metres. Add Δh × Sy × 1000 to recharge millimetres.",
    )
    add_para(
        doc,
        "Availability index. depthScore = clamp(100 − latest/42×100). trendScore = clamp(72 − delta90×38). recScore = clamp(rechargeMm/90×100). index = round(0.45 depth + 0.30 trend + 0.25 rec).",
    )
    add_para(
        doc,
        "Live tick. latest_display = last_series_level + sin(tick×0.7 + i)×0.012×(0.4+stress). Tick increments every 7 s on Overview only.",
    )
    add_para(
        doc,
        "Planner. projected = latest + (delta90/90)×(1−cut/100)×365 − (cut/100)×0.004×(1−0.4×stress)×365, floored at 1 m.",
    )

    doc.add_heading("7. Software architecture", level=1)
    add_para(
        doc,
        "Browser SPA. src/data/stations.js is the 'database'. src/lib/hydrology.js is colour, national aggregates, planner. src/lib/gemini.js is optional NLP. Pages under src/pages. Leaflet map lazy-loaded. "
        "Build: Vite → dist/. Host serves static files. React Router needs rewrite rules (vercel.json, public/_redirects).",
    )
    add_para(
        doc,
        "Folder map: Dashboard.jsx overview; MapPage.jsx filters; StationPage.jsx detail; RechargePage.jsx WTF table; PlannerPage.jsx slider; AboutPage.jsx people and IEEE/CGWB citations.",
    )

    doc.add_heading("8. What is unique (use two of these, not all at once)", level=1)
    for line in [
        "Decision path in one sitting: national map → local DWLR → recharge millimetres → health word → extraction experiment.",
        "WTF method coded on high-frequency-style traces, which matches the synopsis better than a static CGWB PDF screenshot.",
        "Availability index translates three hydrology signals for non-specialists.",
        "Zero-cost static deploy; synopsis AWS was dropped on purpose for a working public URL.",
        "Explicit subset (65 of 5,260) and live-tick labelling — scientific honesty as a feature.",
    ]:
        add_para(doc, "• " + line, space_after=4)

    doc.add_heading("9. Team matrix (from the submitted synopsis)", level=1)
    add_para(doc, "Abhishek Biswal — data analysis and estimation logic (Python/NumPy in the skill matrix; JS recharge engine in this repo).", space_after=4)
    add_para(doc, "Shivam Kumar — GIS and visualisation (Leaflet, Recharts).", space_after=4)
    add_para(doc, "Sushant Kumar — backend and API management (station model, REST-shaped objects; no live Flask host in this build).", space_after=4)
    add_para(doc, "Srushti S Mopagar — architecture and UI (React, Tailwind, hosting).", space_after=4)
    add_para(doc, "Guide: Prof. Prachitha M.", space_after=8)

    doc.add_heading("10. Limitations and ethical notes", level=1)
    add_para(
        doc,
        "Do not frighten a village with a red dot as if it were a government condemnation. Classes are illustrative. Gemini text is not a hydrogeologist’s signed report. "
        "Do not scrape or redistribute restricted DWLR dumps. Keys in frontend can leak; rotate if this repo is public.",
    )

    doc.add_heading("11. Fifty questions the teacher may ask (with answers)", level=1)
    add_para(
        doc,
        "Memorise the short version of each answer. If you freeze, start from question 1, 6, 9, 10, and 17 — those cover 80% of orals.",
    )
    for i, (q, a) in enumerate(QA, 1):
        add_para(doc, f"Q{i}. {q}", bold=True, size=11, space_after=2, color=(15, 42, 40))
        add_para(doc, f"A{i}. {a}", size=11, space_after=10)

    doc.add_heading("12. Quick glossary", level=1)
    glossary = [
        ("Piezometer / observation well", "Well used to measure water level, not primarily to pump."),
        ("Aquifer", "Geologic unit that stores and transmits groundwater."),
        ("Alluvium", "River-laid sand/silt/clay; often higher Sy."),
        ("Hard rock / Deccan basalt / gneiss", "Fractured rocks; low Sy; common in peninsular India."),
        ("Monsoon window in the demo", "June–September style rains in the generator; series ends 24 Aug 2026."),
        ("CARTO / OSM", "Who draws the map background, not who measures wells."),
        ("Gemini", "Google’s large language model; optional sentence writer."),
        ("Vercel / Netlify", "Free hosts for the compiled website."),
    ]
    for k, v in glossary:
        add_para(doc, f"{k}: {v}", space_after=4)

    doc.add_heading("13. References (as in the synopsis)", level=1)
    add_para(
        doc,
        "[1] Central Ground Water Board (CGWB), “National Compilation on Dynamic Ground Water Resources of India,” Ministry of Jal Shakti, Government of India, 2024.",
        space_after=4,
    )
    add_para(
        doc,
        "[2] T. J. Nicholson et al., “Real-time Monitoring of Groundwater Resources,” Journal of Hydrology, 2023.",
        space_after=4,
    )
    add_para(
        doc,
        "[3] IEEE Standard for Sensor Data Middleware and Analytics, IEEE Std 2510-2021.",
        space_after=12,
    )
    add_para(
        doc,
        "— End of guide. Run the app once while reading section 3. Do not commit this .docx to GitHub. —",
        size=10,
        color=(90, 90, 90),
    )

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
